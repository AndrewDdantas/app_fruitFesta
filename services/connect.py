import googlemaps
from datetime import datetime, timedelta
import pandas as pd
import streamlit as st


minha_chave_api = st.secrets['key']
gmaps = googlemaps.Client(key=minha_chave_api)

def converter_segundos(segundos):
    dias = segundos // 86400
    horas = (segundos % 86400) // 3600
    return f'{dias} dia(s) e {horas} hora(s)'

def ajustar_cord(row):

    comma_positions = [i for i, char in enumerate(row) if char == ',']

    if len(comma_positions) == 3 and comma_positions[0] in [2,3,4]:
        row = row.replace(',','.', 1)
        row = row[:comma_positions[2]] + '.' + row[comma_positions[2] + 1:]
        return row.replace(' ','')
    else:
        return row.replace(' ','')

def gerar_endereco(gmaps=gmaps, origem='R. Uirapuru, 316 - Santa Monica, Feira de Santana - BA, 44078-250',waypoints=None):
    matrix = gmaps.distance_matrix(origins=[origem], destinations=waypoints, mode="driving")
    address_distance_list = []
    for destination, way,element in zip(matrix['destination_addresses'], waypoints, matrix['rows'][0]['elements']):
        if element['status'] == 'OK':
            distance = element['distance']['value']  # Distância em metros
        else:
            distance = None  # No caso de falha na busca
        
        address_distance_list.append((destination,way, distance))
    distancia = pd.DataFrame(address_distance_list).drop_duplicates()
    return distancia

def divide_em_blocos(lista, tamanho_bloco):
    origem = '-12.2653916,-38.9302185'
    blocos = []
    lista.insert(0,origem)
    for i in range(0, len(lista), tamanho_bloco):
        bloco = lista[i:i+tamanho_bloco]
        if i != 0:
            bloco.insert(0, blocos[-1][-1])
        blocos.append(bloco)
    bloco.insert(len(blocos[-1]), origem)
    return blocos

def gerar_rota(rotas,gmaps=gmaps):
    informacoes = []
    agora = datetime.now()
    for i in rotas:
        direcoes = gmaps.directions(i[0],
                                    i[-1],
                                    mode="driving",
                                    waypoints=i[1:-1],
                                    optimize_waypoints=True,
                                    departure_time=agora)
        for leg in direcoes[0]['legs']:
            endereco = leg['end_address']
            distancia = leg['distance']['value']
            duracao = leg['duration']['value']
            informacoes.append({'endereço': endereco, 'distância': distancia, 'duração':duracao})
    return informacoes


from docx import Document
from docx.shared import Pt

def adicionar_linha_horizontal():
    o = ''
    i = 1
    while i <= 183:
        o += "_" 
        i += 1
    return o

# Função para substituir as chaves no texto
def substituir_chaves(paragraphs, substitutions):
    for paragraph in paragraphs:
        for key, value in substitutions.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
                
def criar_romaneio(merge,volumes, tempo, km):
# Abrir o documento existente
    doc = Document('./romaneio_modelo.docx')

    # Dicionário com as chaves e os valores para substituição
    substituicoes = {
        '<<tt_rota>>': km,
        '<<pc>>': f'{volumes} peças',
        '<<tm>>': tempo
    }

    # Substituir as chaves no documento
    for paragraph in doc.paragraphs:
        substituir_chaves([paragraph], substituicoes)

    # Substituir as chaves em tabelas (se houver)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_chaves(cell.paragraphs, substituicoes)

    # Lista de dados fictícios


    # Adicionar parágrafos para cada cliente
    for index, row in merge.iterrows():
        # Nome do cliente
        doc.add_paragraph(f"Pedido: {row['OV']} | Nota: {row['Nota']}| Cliente: {row['ClienteDesc']} | Endereço: {row['endereço']}\nAssinatura:{adicionar_linha_horizontal()}").paragraph_format.space_after = Pt(10)

    return doc
